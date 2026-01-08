import os
import pandas as pd
import json
import logging
import sqlite3
from datetime import datetime
from typing import Dict
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langchain_ollama import OllamaEmbeddings
from langchain_chroma import Chroma
from langchain_core.prompts import ChatPromptTemplate
from langchain_ollama.llms import OllamaLLM
from langchain_core.output_parsers import StrOutputParser

# Define Folders
load_dotenv()
DB_NAME = "gaia.db"
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "./uploaded_pdfs") # Documents Dir
PERSIST_DIR = os.getenv("PERSIST_DIR", "./chroma_store") # Vector Data Dir
REPORTS_DIR = os.getenv("REPORTS_DIR", "./reports") # Reports Dir
os.makedirs(REPORTS_DIR, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Embedding & LLM Models
embeddings = OllamaEmbeddings(model="mxbai-embed-large")
llm = OllamaLLM(model="qwen3-vl:235b-cloud", base_url="http://localhost:11434")

# Initialize DB
def init_db():
  """Initializes the SQLite database with the sessions table."""
  con = sqlite3.connect(DB_NAME)
  c = con.cursor()
  # Table: Roles (Categories)
  c.execute('''CREATE TABLE IF NOT EXISTS roles (
      role_id TEXT PRIMARY KEY,
      role_name TEXT
  )''')

  # Table: Scenarios (The Content Pairs)
  # Linked to Role. Stores the Prompts.
  c.execute('''CREATE TABLE IF NOT EXISTS scenarios (
      scenario_id TEXT PRIMARY KEY,
      role_id TEXT,
      topic TEXT,
      mentor_persona TEXT,
      simulation_persona TEXT,
      scenario_details TEXT,
      FOREIGN KEY(role_id) REFERENCES roles(role_id)
  )''')

  # Table: Grading Rubrics (Prompt Parameters)
  # Stores the criteria list for the AI.
  c.execute('''CREATE TABLE IF NOT EXISTS grading_rubrics (
      rubric_id INTEGER PRIMARY KEY AUTOINCREMENT,
      scenario_id TEXT,
      criteria TEXT,
      description TEXT,
      FOREIGN KEY(scenario_id) REFERENCES scenarios(scenario_id)
  )''')

  # Table: Sessions (Header)
  # Replaced 'user_id' with 'trainee_name' for simple tracking.
  c.execute('''CREATE TABLE IF NOT EXISTS sessions (
      session_id TEXT PRIMARY KEY,
      trainee_name TEXT,
      scenario_id TEXT,
      date TEXT,
      total_score INTEGER,
      readiness TEXT,
      chat_log TEXT,
      report_path TEXT,
      FOREIGN KEY(scenario_id) REFERENCES scenarios(scenario_id)
  )''')

  # Table: Session Grades (Detail)
  # Stores specific scores/feedback per criteria.
  c.execute('''CREATE TABLE IF NOT EXISTS session_grades (
      grade_id INTEGER PRIMARY KEY AUTOINCREMENT,
      session_id TEXT,
      criteria TEXT,
      score INTEGER,
      evidence TEXT,
      feedback TEXT,
      FOREIGN KEY(session_id) REFERENCES sessions(session_id)
  )''')

  con.commit()

  # CHECK DATA EXISTENCE
  c.execute("SELECT count(*) FROM roles")
  data_count = c.fetchone()[0]
  
  con.close()

  # Only seed if roles table is empty
  if data_count == 0:
      logger.info("Database is empty. Seeding initial data...")
      seed_db()
  else:
      logger.info("Database already contains data. Skipping seed.")

def seed_db():
    """
    Populates the DB with initial Role/Scenario data AND Dummy Sessions
    to ensure the Dashboard functions correctly out-of-the-box.
    """
    con = sqlite3.connect(DB_NAME)
    c = con.cursor()

    # Check if roles exist to avoid duplicates
    c.execute("SELECT count(*) FROM roles")
    if c.fetchone()[0] == 0:
        print("Seeding Core Data (Roles, Scenarios, Rubrics)...")
        
        # 1. Insert Roles
        roles = [
            ("CS", "Customer Service"),
            ("TELLER", "Bank Teller")
        ]
        c.executemany("INSERT INTO roles VALUES (?,?)", roles)

        # 2. Insert Scenarios
        scenarios = [
            (
                "TELLER_CASH", "TELLER", 
                "Penanganan Uang Meragukan (Counterfeit) pada Nasabah Prioritas",
                "Anda adalah Senior Head Teller bernama 'Pak Teguh'. Anda adalah perwujudan dari 'Zero Tolerance Policy'.", 
                "Anda adalah 'Bapak Hartono', nasabah Prioritas (Solitaire) pemilik jaringan ritel terbesar di kota ini.", 
                "Bapak Hartono menyetor Rp 200 Juta tunai hasil penjualan toko. Mesin hitung menolak (reject) 2 lembar pecahan Rp 100.000."
            ),
            (
                "CS_COMPLAINT", "CS", 
                "Handling Panic Customer: Indikasi Social Engineering (Fraud)",
                "Anda adalah Service Quality Manager bernama 'Ibu Sari'. Anda fokus pada 'Customer Journey' dan 'Empathy'.",
                "Anda adalah 'Ibu Lina', seorang pengusaha katering. Anda baru saja menerima telepon yang mengaku dari pihak bank.", 
                "Nasabah datang dengan histeris karena saldonya terkuras setelah mengklik file .APK undangan pernikahan (Phishing)."
            )
        ]
        c.executemany("INSERT INTO scenarios VALUES (?,?,?,?,?,?)", scenarios)

        # 3. Insert Rubrics
        rubrics = [
            ("TELLER_CASH", "Sikap Profesional & Tenang", "Trainee tidak boleh terlihat gugup."),
            ("TELLER_CASH", "Pemilihan Kata (Euphemism)", "DILARANG menggunakan kata 'PALSU' sebelum verifikasi."),
            ("CS_COMPLAINT", "Immediate Security Action", "Langkah pertama Trainee HARUS melakukan pemblokiran."),
            ("CS_COMPLAINT", "Empati Tanpa Menjanjikan", "Mengucapkan keprihatinan mendalam tanpa janji palsu.")
        ]
        c.executemany("INSERT INTO grading_rubrics (scenario_id, criteria, description) VALUES (?,?,?)", rubrics)

    # 4. Insert Dummy Sessions (For Dashboard Visualization)
    c.execute("SELECT count(*) FROM sessions")
    if c.fetchone()[0] == 0:
        print("Seeding Dummy Session Data...")
        
        # Format: session_id, trainee_name, scenario_id, date, total_score, readiness, chat_log
        dummy_sessions = [
            ("SES-101", "Andi Saputra", "TELLER_CASH", "2024-10-01 09:30", 88, "SIAP TERJUN", "System: Welcome..."),
            ("SES-102", "Budi Santoso", "CS_COMPLAINT", "2024-10-02 10:15", 55, "BELUM SIAP", "System: Welcome..."),
            ("SES-103", "Citra Lestari", "TELLER_CASH", "2024-10-03 11:00", 92, "SIAP TERJUN", "System: Welcome..."),
            ("SES-104", "Dewi Persik", "CS_COMPLAINT", "2024-10-04 13:45", 76, "BUTUH LATIHAN", "System: Welcome..."),
            ("SES-105", "Eko Patrio", "TELLER_CASH", "2024-10-05 14:30", 65, "BUTUH LATIHAN", "System: Welcome..."),
            ("SES-106", "Fajar Hadi", "CS_COMPLAINT", "2024-10-06 15:00", 40, "BELUM SIAP", "System: Welcome..."),
            ("SES-107", "Gita Gutawa", "TELLER_CASH", "2024-10-07 08:30", 95, "SIAP TERJUN", "System: Welcome..."),
            ("SES-108", "Hesti Purwadinata", "CS_COMPLAINT", "2024-10-07 09:00", 82, "SIAP TERJUN", "System: Welcome..."),
            ("SES-109", "Indra Bekti", "TELLER_CASH", "2024-10-08 10:45", 70, "BUTUH LATIHAN", "System: Welcome..."),
            ("SES-110", "Joko Anwar", "CS_COMPLAINT", "2024-10-08 11:30", 89, "SIAP TERJUN", "System: Welcome...")
        ]
        c.executemany("INSERT INTO sessions VALUES (?,?,?,?,?,?,?)", dummy_sessions)

        # 5. Insert Dummy Grades (Linked to Sessions)
        # Format: session_id, criteria, score, evidence, feedback
        dummy_grades = [
            ("SES-101", "Sikap Profesional & Tenang", 90, "Nasabah marah, trainee tetap senyum", "Good job maintaining composure."),
            ("SES-101", "Pemilihan Kata (Euphemism)", 85, "Menggunakan istilah 'diragukan'", "Tepat sekali."),
            ("SES-102", "Immediate Security Action", 20, "Hanya mendengarkan curhat nasabah", "Fatal: Lupa blokir rekening."),
            ("SES-102", "Empati Tanpa Menjanjikan", 90, "Saya turut prihatin bu", "Empati bagus."),
        ]
        c.executemany("INSERT INTO session_grades (session_id, criteria, score, evidence, feedback) VALUES (?,?,?,?,?)", dummy_grades)

    con.commit()
    con.close()
    print("Database seeding complete.")

# Logger
def setup_logger(name="gaia"):

    logger = logging.getLogger(name)
    logger.setLevel(logging.DEBUG)

    # Console Handler
    ch = logging.StreamHandler()
    ch.setLevel(logging.DEBUG)

    # Formatter
    formatter = logging.Formatter("[%(asctime)s] [%(levelname)s] - %(message)s ")
    ch.setFormatter(formatter)

    # Duplicate Log Handler
    if not logger.hasHandlers():
        logger.addHandler(ch)

    return logger

logger = setup_logger()

def format_chat_history(messages: list) -> str:
  formatted_history = ""
  for msg in messages:
    role = msg["role"].upper()
    content = msg["content"]
    formatted_history += f"{role}: {content}\n"
  return formatted_history

# Load the vectors
def load_vectors():
    vectorstore = Chroma(
        persist_directory=PERSIST_DIR,
        embedding_function=embeddings,
        collection_name="knowledge_base"
    )

    return vectorstore

# Retrieve the chunks
def get_retriever(vectorstore, k=3):
    return vectorstore.as_retriever(search_kwargs={"k": k})

# NOTE: NOT USED
DUMMY_DB = [  
  {
    "role_name": "TELLER_CASH",
    "topic": "Penanganan Uang Meragukan (Counterfeit) pada Nasabah Prioritas",
    "mentor_persona": "Anda adalah Senior Head Teller bernama 'Pak Teguh'. Anda adalah perwujudan dari 'Zero Tolerance Policy'. Bagi Anda, melindungi bank dari risiko operasional dan reputasi adalah segalanya. Gaya bicara Anda tegas, instruktif, dan selalu menekankan pada prinsip 3D (Dilihat, Diraba, Diterawang) serta pelaporan berjenjang.",
    "simulation_persona_text": "Anda adalah 'Bapak Hartono', nasabah Prioritas (Solitaire) pemilik jaringan ritel terbesar di kota ini. Anda sangat sibuk dan merasa prosedur bank yang berbelit-belit hanya membuang waktu Anda. Anda tipe orang yang biasa dilayani VIP. Jika ada hambatan, Anda cenderung langsung menelpon Pimpinan Cabang daripada berdebat dengan staf.",
    "scenario_details_text": "Bapak Hartono menyetor Rp 200 Juta tunai hasil penjualan toko. Mesin hitung menolak (reject) 2 lembar pecahan Rp 100.000. Saat diperiksa manual, kertas terasa halus dan benang pengaman tidak berubah warna (Indikasi Palsu). Konflik: Bapak Hartono tersinggung uangnya diragukan, mengklaim itu uang dari ATM bank ini juga, dan mengancam akan memindahkan seluruh saldo depositonya jika Anda mempermasalahkan 'uang receh' 200 ribu tersebut.",
    "success_criteria": [
      {
        "criteria": "Sikap Profesional & Tenang",
        "description": "Trainee tidak boleh terlihat gugup atau terintimidasi oleh status nasabah. Tetap melakukan 3D (Dilihat, Diraba, Diterawang) secara transparan di hadapan nasabah."
      },
      {
        "criteria": "Pemilihan Kata (Euphemism)",
        "description": "DILARANG menggunakan kata 'PALSU' sebelum verifikasi final. Gunakan frasa: 'Maaf Bapak, ada beberapa lembar yang tidak lolos sensor mesin dan perlu kami verifikasi manual'."
      },
      {
        "criteria": "Handling Objection (Threat)",
        "description": "Saat nasabah mengancam memindahkan dana, Trainee tetap tenang dan menjelaskan bahwa prosedur ini justru untuk melindungi nasabah dari peredaran uang yang tidak layak, bukan menuduh nasabah."
      },
      {
        "criteria": "Prosedur Penahanan",
        "description": "Menjelaskan aturan Bank Indonesia bahwa fisik uang harus ditahan untuk dikirim ke BI (Klarifikasi), dan memberikan tanda terima penahanan uang kepada nasabah."
      }
    ]
  },
  {
    "role_name": "CS_COMPLAINT",
    "topic": "Handling Panic Customer: Indikasi Social Engineering (Fraud)",
    "mentor_persona": "Anda adalah Service Quality Manager bernama 'Ibu Sari'. Anda fokus pada 'Customer Journey' dan 'Empathy'. Motto Anda: 'Nasabah mungkin salah karena memberikan OTP, tapi mereka adalah korban kejahatan yang sedang panik. Jangan hakimi mereka, tapi lindungi aset mereka.",
    "simulation_persona_text": "Anda adalah 'Ibu Lina', seorang pengusaha katering. Anda baru saja menerima telepon yang mengaku dari pihak bank, lalu saldo rekening Anda berkurang Rp 15 Juta. Anda sangat panik, marah, menangis, dan menyalahkan sistem keamanan bank yang lemah. Anda menuntut uang kembali detik ini juga.",
    "scenario_details_text": "Nasabah datang dengan histeris karena saldonya terkuras setelah mengklik file .APK undangan pernikahan (Phishing). Nasabah tidak sadar bahwa itu kesalahannya dan menuntut Bank bertanggung jawab. Trainee harus melakukan pemblokiran darurat, menenangkan nasabah, menggali kronologi tanpa menghakimi, namun tetap tegas menjelaskan bahwa proses pengembalian dana membutuhkan investigasi dan tidak bisa instan.",
    "success_criteria": [
      {
        "criteria": "Immediate Security Action",
        "description": "Langkah pertama Trainee HARUS melakukan pemblokiran akun/kartu untuk mencegah kerugian lebih lanjut sebelum mendengarkan cerita panjang lebar."
      },
      {
        "criteria": "Empati Tanpa Menjanjikan (No False Promise)",
        "description": "Mengucapkan keprihatinan mendalam ('Saya turut prihatin atas kejadian ini Bu'), NAMUN tidak boleh menjanjikan uang pasti kembali."
      },
      {
        "criteria": "Investigasi Kronologis (Fact Finding)",
        "description": "Menggali data sensitif dengan hati-hati: 'Apakah Ibu sempat memberikan kode OTP atau mengklik tautan/aplikasi di luar PlayStore?'"
      },
      {
        "criteria": "Edukasi & Ekspektasi",
        "description": "Menjelaskan prosedur investigasi (SLA kerja), pembuatan laporan kepolisian, dan mengedukasi nasabah tentang bahaya file APK/Phishing agar tidak terulang."
      }
    ]
  }
]

def fetch_roleplay_data(scenario_id: str) -> Dict:
    # Fetching the data
    role_data = get_scenario_config(scenario_id)
    if not role_data:
        raise ValueError(f"Scenario ID {scenario_id} not found.")
    return role_data

def get_scenario_config(scenario_id):
  """
  Retrieves the full configuration (Personas + Rubrics) for the AI Engine.
  Replacing 'fetch_roleplay_data' logic.
  """
  con = sqlite3.connect(DB_NAME)
  con.row_factory = sqlite3.Row
  c = con.cursor()

  # 1. Get Main Scenario Data
  c.execute("SELECT * FROM scenarios WHERE scenario_id = ?", (scenario_id,))
  row = c.fetchone()

  if not row:
    con.close()
    return None
  
  scenario_data = dict(row)

  # 2. Get Rubrics List
  c.execute("SELECT criteria, description FROM grading_rubrics WHERE scenario_id = ?", (scenario_id,))
  rubrics = [dict(r) for r in c.fetchall()]

  con.close()

  return {
      "role_name": scenario_data["scenario_id"],
      "topic": scenario_data["topic"],
      "mentor_persona": scenario_data["mentor_persona"],
      "simulation_persona_text": scenario_data["simulation_persona"],
      "scenario_details_text": scenario_data["scenario_details"],
      "success_criteria": rubrics
  }

def save_full_session(session_data, grade_list):
    """
    Transactional Save: Stores the Session Header AND the Detailed Grades.
    """
    con = sqlite3.connect(DB_NAME)
    c = con.cursor()

    # 1. Save Session Header
    try:
        c.execute('''INSERT INTO sessions
            (session_id, trainee_name, scenario_id, date, total_score, readiness, chat_log, report_path)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
            (
                session_data['session_id'],     
                session_data['trainee_name'],     
                session_data['scenario_id'],     
                session_data['date'],     
                session_data['total_score'],
                session_data['readiness'],     
                str(session_data['chat_log']),
                session_data.get('report_path', '')
            )
        )

        # 2. Save Detailed Grades
        for grade in grade_list:
            c.execute('''INSERT INTO session_grades
                (session_id, criteria, score, evidence, feedback)
                VALUES (?, ?, ?, ?, ?)''',
                (
                    session_data['session_id'],
                    grade['criteria'],
                    grade['score'],
                    grade['evidence'],
                    grade['feedback'],
                )
            )
        con.commit()
    except Exception as e:
        print(f"Error saving session: {e}")
        con.rollback()
        raise
    finally:
        con.close()

def fetch_all_sessions():
    """
    Returns all session headers for the Dashboard Table.
    Joins with Scenarios to get Topic names.
    """
    con = sqlite3.connect(DB_NAME)
    query = '''
        SELECT
            s.session_id,
            s.trainee_name,
            sc.topic as Role,
            s.date,
            s.total_score as Score,
            s.readiness,
            sc.role_id,
            s.report_path
        FROM sessions s
        JOIN scenarios sc ON s.scenario_id = sc.scenario_id
        ORDER BY s.date DESC
    '''
    try:
        df = pd.read_sql_query(query, con)

        # Data Cleaning & Feature Engineering
        if not df.empty:
            # 1. Create Status Column based on 'Score'
            df['Status'] = df['Score'].apply(lambda x: "Passed" if x >= 80 else "Failed")

            # 2. Handle data formatting
            df['date'] = pd.to_datetime(df['date'])
            # 3. Duration Placeholder
            df['Duration (Mins)'] = 15
    except ImportError:
        # Fallback if pandas is not installed (returns list of dicts)
        con.row_factory = sqlite3.Row
        c = con.cursor()
        c.execute(query)
        df = [dict(row) for row in c.fetchall()]
    finally:
        con.close()

    return df

def fetch_session_details(session_id):
    """
    Fetches the Grade Breakdown for a specific session (Drill Down)
    """
    con = sqlite3.connect(DB_NAME)
    con.row_factory = sqlite3.Row
    c = con.cursor()

    c.execute("SELECT * FROM session_grades WHERE session_id = ?", (session_id,))
    grades = [dict(row) for row in c.fetchall()]

    con.close()
    return grades

def build_system_prompt(phase: str, data: dict) -> str:
    """
    Constructs the System Prompt dynamically based on the current Phase 
    and the Topic Configuration (data) retrieved from the Database.
    """

    # Variables
    mentor_persona = data.get("mentor_persona")
    topic = data.get("topic")
    simulation_persona = data.get("simulation_persona_text")
    scenario_details = data.get("scenario_details_text")
    success_criteria = json.dumps(data.get("success_criteria"), indent=2)

    # ---------------------------------------------------------
    # PHASE 1 & 2: GREETING & FLOW EXPLANATION
    # ---------------------------------------------------------
    if phase == "GREETING":
        return f"""
        ### SYSTEM ROLE & IDENTITY
        You are {mentor_persona}.
        Your goal is to guide a trainee through a learning session about: {topic}.

        ### CURRENT OBJECTIVE: ESTABLISH RAPPORT & SET EXPECTATIONS
        You are currently in **PHASE 1 (Opening)** and **PHASE 2 (Flow Explanation)**.

        **INSTRUCTIONS:**
        1. **Greet:** Welcome the trainee professionally. State your name and role clearly.
        2. **Explain the Case:** Briefly explain that the roleplay session case about {scenario_details}
        3. **Transition (CRITICAL):** - Ask if they are ready to begin.
          - If they say yes, **instruct them explicitly** to click the button **'🎓 Start Tutoring Session'** below to proceed.
        """
    
    # ---------------------------------------------------------
    # PHASE 3: TUTORING SESSION
    # ---------------------------------------------------------
    elif phase == "TUTORING":
        return f"""
        ### SYSTEM ROLE & IDENTITY
        You are {mentor_persona}.
        Topic: {topic}.

        ### CURRENT OBJECTIVE: PHASE 3 - TUTORING SESSION (ASSESSMENT)
        **Goal:** Gauge the trainee's understanding of the topic using the Knowledge Base.

        **INSTRUCTIONS:**
        1. **Ask Questions:** Ask 2-3 specific, challenging questions about {topic} to test their theoretical knowledge.
        2. **Answer Questions:** Allow the trainee to ask questions back. Answer them clearly using the **[CONTEXT/KNOWLEDGE BASE]** provided below.
        3. **Tone:** Be helpful, educational, and supportive.
        4. **Correction:** If they answer incorrectly, correct them gently using facts from the Knowledge Base.
        5. **Transition (CRITICAL):** - Once you are satisfied with their understanding (or after 3 interactions), say: "We are now ready for the simulation."
          - **Instruct them explicitly** to click the button **'🚀 Start Roleplay Simulation'** below to enter the scenario.

        **CONSTRAINT:** Do NOT start the Roleplay simulation yet. Stay in the Tutoring phase.
        """
    
    # ---------------------------------------------------------
    # PHASE 4: ROLEPLAY SIMULATION
    # ---------------------------------------------------------
    elif phase == "ROLEPLAY":
        return f"""
        ### SYSTEM MODE SWITCH: SIMULATION (STRICT MODE)
        **CRITICAL:** STOP being the Mentor. BECOME the following persona:
        {simulation_persona}

        ### SCENARIO CONTEXT
        {scenario_details}

        ### CURRENT OBJECTIVE: PHASE 4 - ROLEPLAY SIMULATION
        **Goal:** Test the trainee's application of skills in a realistic scenario.

        **CONSTRAINTS & BEHAVIOR:**
        1. **Do NOT give advice.** You are the customer/counterpart, not the teacher.
        2. **Do NOT break character.** Even if the trainee asks for help, reply IN CHARACTER (e.g., "I don't care about your manual, I want this fixed!").
        3. **Emotional Reaction:** React realistically. Get happier if they handle it well, get angrier/more frustrated if they stick to scripts that don't help.
        4. **Transition (CRITICAL):** - Continue until the problem is solved or the trainee gives up.
          - When the scene ends, BREAK CHARACTER immediately.
          - Say: "SIMULATION ENDED."
          - **Instruct them explicitly** to click the button **'🏁 Finish & Grade'** below to see their score.
        """
    
    # ---------------------------------------------------------
    # PHASE 5: GRADING & SUMMARY
    # ---------------------------------------------------------
    elif phase == "GRADING":      
        return f"""
        ### SYSTEM MODE SWITCH: AUDITOR
        Revert to your original persona: {mentor_persona}.

        ### CURRENT OBJECTIVE: PHASE 5 - SUMMARY & SCORING
        **Goal:** Provide detailed feedback on the simulation and assess real-world readiness.

        **INSTRUCTIONS:**
        Review the conversation history (focusing on the Roleplay phase).
        1. Generate a Grading Table based on the rubric.
        2. Determine the **Readiness Level (Tingkat Kesiapan)** of the trainee to face this scenario in real life.

        **GRADING RUBRIC:**
        {success_criteria}

        **OUTPUT FORMAT:**
        Generate a Markdown table with exactly these columns:
        | Criteria | Evidence (Quote) | Feedback | Score (0-100) |

        **Readiness Level:**
        Analyze the total performance and assign a status:
        - **SIAP TERJUN (Ready):** If the trainee handled critical constraints well and scored high (>80).
        - **BUTUH LATIHAN (Needs Practice):** If they missed some protocols but understood the basics (60-79).
        - **BELUM SIAP (Not Ready):** If they failed critical criteria or broke character (<60).

        *Display it clearly below the table like this:*
        > **Status Kesiapan:** [SIAP TERJUN / BUTUH LATIHAN / BELUM SIAP]
        > **Kesimpulan:** [One sentence explaining why]

        **CLOSING:**
        Offer a final encouraging remark to the trainee.
        """

    return mentor_persona

# Chain Query
def query_chain(retriever, llm, user_input: str, role_id: str, current_phase: str, chat_history: list):
    """
    Orchestrates the entire flow: Data -> Prompt -> RAG -> LLM
    """

    try:
        logger.info(f"--- Starting Chain: {role_id} | Phase: {current_phase} ---")

        # Fetch Data from DB
        role_data = fetch_roleplay_data(role_id)

        # Build Dynamic System Prompt
        system_instructions = build_system_prompt(current_phase, role_data)

        # Update History
        history_text = format_chat_history(chat_history)

        # Optimization: Only retrieve docs in 'TUTORING'. In 'ROLEPLAY', context is the scenario.
        if current_phase == "TUTORING":
            knowledge_base_content = retriever.invoke(user_input)
        elif current_phase == "GREETING":
            knowledge_base_content = "Session Initiated."
        else:
            knowledge_base_content = "Refer to Scenario Details in System Prompt."

        # Prompt Template
        template = """
        {role_instruction}

        [CONTEXT/KNOWLEDGE BASE]: {knowledgeBase}

        [CHAT HISTORY]: {history}

        [USER INPUT]: {question}

        [STRICT GUIDELINES]
        1. You are strictly prohibited from answering questions that are NOT related to the [CONTEXT/KNOWLEDGE BASE] or [CHAT HISTORY].
        2. If the user asks a question outside of the provided scope, you must politely decline and state that you can only answer questions related to the specific context provided.
        3. Do not use outside knowledge or general training data to answer unrelated queries.
        4. ALWAYS RESPOND IN BAHASA INDONESIA
        """

        # Build the Chain
        prompt = ChatPromptTemplate.from_template(template)
        chain = prompt | llm | StrOutputParser()

        # Invoke
        result = chain.invoke({"role_instruction": system_instructions, "knowledgeBase": knowledge_base_content, "history": history_text, "question": user_input})
        return result
    except Exception as e:
        logger.exception("Error querying the chain")
        raise

def create_individual_report(session_data, grades_list, chat_history, llm):
    """
    Generates a full performance report.
    - Uses the LLM to generate a qualitative 'Readiness Assessment'.
    - Creates a Word Document with: Meta Data -> AI Insight -> Grading Matrix -> Transcript.
    """

    doc = Document()

    try:
        # ==========================================
        # 1. GENERATE AI INSIGHT (The New Logic)
        # ==========================================
        template = f"""
        You are a Senior Banking Training Mentor. 
        Write a formal "Performance Review & Readiness Assessment" (2-3 paragraphs) for a trainee who just completed the "{role_played}" simulation.

        **Session Data:**
        - Final Score: {score}/100
        - Grading Details: {grading_summary}

        **Instructions:**
        1. Summarize their key strengths based on the high scores.
        2. Point out specific critical errors (if any) based on low scores/feedback.
        3. Provide a clear recommendation: Are they ready for the real job? If not, what specific SOPs must they re-read?
        4. Tone: Professional, constructive, and encouraging.
        """

        # Prepare the context for the AI
        grading_summary = json.dumps(grades_list, indent=2)
        role_played = session_data.get('scenario_id', 'unknown')
        score = session_data.get('total_score', 0)

        prompt = ChatPromptTemplate.from_template(template)
        chain = prompt | llm | StrOutputParser()

        result = chain.invoke({"role_played": role_played, "score": score, "grading_summary": grading_summary})

        # --- HEADER ---
        header = doc.add_heading(f"Trainee Performance Report", 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        # --- SECTION 1: SESSION METADATA ---
        doc.add_heading("1. Session Details", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Column 1
        cell1 = table.cell(0, 0)
        p = cell1.paragraphs[0]
        p.add_run("Trainee Name: ").bold = True
        p.add_run(f"{session_data.get('trainee_name', 'Guest')}\n")
        p.add_run("Role: ").bold = True
        p.add_run(f"{role_played}\n")
        p.add_run("Session ID: ").bold = True
        p.add_run(f"{session_data.get('session_id', 'N/A')}")

        # Column 2
        cell2 = table.cell(0, 1)
        p = cell2.paragraphs[0]
        p.add_run("Total Score: ").bold = True
        
        # Dynamic Color for Score
        score_run = p.add_run(f"{score}/100")
        score_run.bold = True
        score_run.font.size = Pt(14)
        if score >= 80:
            score_run.font.color.rgb = RGBColor(0, 128, 0) # Green
        elif score >= 60:
            score_run.font.color.rgb = RGBColor(255, 165, 0) # Orange
        else:
            score_run.font.color.rgb = RGBColor(255, 0, 0) # Red
            
        p.add_run(f"\nStatus: {session_data.get('readiness', 'Unknown')}")

        doc.add_paragraph() # Spacer

        # --- SECTION 2: AI MENTOR INSIGHT (The "Executive Summary" equivalent) ---
        doc.add_heading("2. Mentor Readiness Assessment", level=1)
        insight_p = doc.add_paragraph(result)
        insight_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # --- SECTION 3: GRADING MATRIX ---
        doc.add_heading("3. Grading Breakdown", level=1)
        
        g_table = doc.add_table(rows=1, cols=4)
        g_table.style = 'Table Grid'
        
        # Header
        hdr_cells = g_table.rows[0].cells
        headers = ["Criteria", "Score", "Evidence", "Feedback"]
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            
        # Rows
        for grade in grades_list:
            row_cells = g_table.add_row().cells
            row_cells[0].text = str(grade.get('criteria', ''))
            row_cells[1].text = str(grade.get('score', ''))
            row_cells[2].text = str(grade.get('evidence', ''))
            row_cells[3].text = str(grade.get('feedback', ''))

        # --- SECTION 4: TRANSCRIPT (Optional/Appendix) ---
        doc.add_heading("Appendix: Chat Transcript", level=1)
        
        for msg in chat_history:
            role = msg['role'].upper()
            content = msg['content']
            
            # Skip internal triggers
            if "[SYSTEM_TRIGGER" in content:
                continue
                
            p = doc.add_paragraph()
            role_run = p.add_run(f"{role}: ")
            role_run.bold = True
            role_run.font.size = Pt(9)
            p.add_run(content).font.size = Pt(9)

        # --- SAVE ---
        filename = f"{REPORTS_DIR}/{session_data['session_id']}_{session_data.get('trainee_name', 'User').replace(' ', '_')}.docx"
        doc.save(filename)
        return filename
        
    except Exception as e:
        logger.exception("Error generating Individual Report Results")
        raise

def create_executive_summary(overall_stats, data_summary, llm):
    """
    Generates a Word doc for the PIC with aggregate insights.
    """
    doc = Document()
    # Generate AI Summary
    try:
        template = f"""
        You are a Senior Training Consultant. Analyze the following training data CSV:
        {data_summary}
        
        Write an Executive Summary (3-5 paragraphs) covering:
        1. **Overall Performance Trends**: Are trainees generally passing?
        2. **Common Weaknesses**: Identify roles or metrics with low scores.
        3. **Strategic Recommendations**: What should the training team focus on next week?

        Use Bahasa Indonesia to generate the executive summary
        """

        prompt = ChatPromptTemplate.from_template(template)
        chain = prompt | llm | StrOutputParser()

        result = chain.invoke({"data_summary": data_summary})

        # --- HEADER ---
        title = doc.add_heading("Executive Training Summary", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d')}")

        # --- SECTION 1: HIGH LEVEL STATS ---
        doc.add_heading("1. Training Overview", level=1)
        p = doc.add_paragraph()
        p.add_run(f"Total Sessions: {overall_stats['total_sessions']}\n")
        p.add_run(f"Average Score: {overall_stats['avg_score']:.2f}\n")
        p.add_run(f"Pass Rate: {overall_stats['pass_rate']}%")

        # --- SECTION 2: AI STRATEGIC ANALYSIS ---
        doc.add_heading("2. AI Strategic Analysis", level=1)
        doc.add_paragraph(result) 

        # --- SAVE ---
        filename = f"{REPORTS_DIR}/Executive_Summary_{datetime.now().strftime('%Y%m%d')}.docx"
        doc.save(filename)
        return filename
        
    except Exception as e:
        logger.exception("Error generating Executive Summary")
        raise